"""Scrittura file di output (TXT, DOCX, HTML) con gestione dei tag markup di Surya.

Tag supportati (prodotti da Surya):
  <i>…</i>     → corsivo
  <sup>N</sup>  → apice (numero di nota)

Per .txt: <sup> → apice Unicode (¹²³…), <i> → testo nudo.
Per .docx: formattazione Word nativa (corsivo e apice reali).
Per .docx (Surya 0.20): rendering strutturato con heading, tabelle, caption, note.
Per .html (Surya 0.20): HTML completo con CSS pronto per la visualizzazione.
"""

import os
import re
import sys
from html import unescape
from html.parser import HTMLParser

import docx

from app.i18n import _

# Mappa cifre ASCII → apici Unicode (U+2070 … U+2079)
_SUPERSCRIPT_MAP = str.maketrans("0123456789", "⁰¹²³⁴⁵⁶⁷⁸⁹")

# Pattern per i tag di primo livello (non-greedy, DOTALL per contenuti multiriga)
_TAG_RE = re.compile(r"(<i>.*?</i>|<sup>\d+</sup>|<math>.*?</math>)", re.DOTALL)
# Pattern per <sup> eventualmente annidato dentro <i>
_SUP_RE = re.compile(r"(<sup>\d+</sup>)")

# Stili Word per etichette Surya 0.20.
# NB: Surya 2 usa label canoniche in camelCase (SectionHeader, PageHeader…),
# non i vecchi nomi con trattino. Surya 2 non ha una label "Title": il livello
# di intestazione più alto è "SectionHeader".
_BLOCK_STYLE = {
    "SectionHeader": "Heading 1",
    "Caption": "Caption",
    "PageHeader": "Header",
    "PageFooter": "Footer",
}

# Blocchi puramente visivi: nel DOCX diventano un segnaposto testuale, così gli
# utenti non vedenti sanno dove si trovavano le immagini nel layout originale.
_IMAGE_PLACEHOLDER = {
    "Picture": "[Image]",
    "Figure": "[Figure]",
}

# Etichette da saltare nel DOCX: solo le pagine vuote (nessun contenuto utile).
_SKIP_DOCX_LABELS = frozenset({"BlankPage"})


def strip_markup(text: str) -> str:
    """Rimuove tutti i tag HTML-like (usato per la visualizzazione in anteprima)."""
    text = text.replace("\f", "\n\n")
    return re.sub(r"<[^>]+>", "", text)


def _convert_markup(text: str) -> str:
    """Converte i tag in equivalenti plain-text/Unicode (per .txt).

    - <sup>N</sup>  → apice Unicode (¹, ², ¹²…)
    - <i>…</i>     → solo il contenuto
    - <br>          → newline
    - altri tag     → rimossi
    """
    text = text.replace("\f", "\n\n")
    text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(
        r"<sup>(\d+)</sup>",
        lambda m: m.group(1).translate(_SUPERSCRIPT_MAP),
        text,
    )
    text = re.sub(r"<i>(.*?)</i>", r"\1", text, flags=re.DOTALL)
    text = re.sub(r"<math>(.*?)</math>", r"\1", text, flags=re.DOTALL)
    text = re.sub(r"<[^>]+>", "", text)
    return text


def _add_markup_runs(para, text: str) -> None:
    """Aggiunge runs con formattazione Word nativa al paragrafo.

    Gestisce:
    - testo normale → run standard
    - <i>…</i>     → run corsivo (con eventuale <sup> annidato)
    - <sup>N</sup>  → run apice
    - \\n interno   → a capo morbido (Shift+Enter in Word)
    """

    def _flush(content: str, italic: bool = False) -> None:
        """Aggiunge i run per `content`, gestendo <sup> interni e \\n."""
        for seg in _SUP_RE.split(content):
            if not seg:
                continue
            is_sup = seg.startswith("<sup>") and seg.endswith("</sup>")
            if is_sup:
                inner = seg[5:-6]
            else:
                inner = re.sub(r'<br\s*/?>', '\n', seg, flags=re.IGNORECASE)
                inner = re.sub(r'<[^>]+>', '', inner)
            lines = inner.split("\n")
            for j, line in enumerate(lines):
                if j > 0:
                    para.add_run().add_break()
                if line:
                    run = para.add_run(line)
                    if italic:
                        run.italic = True
                    if is_sup:
                        run.font.superscript = True

    for segment in _TAG_RE.split(text):
        if not segment:
            continue
        if segment.startswith("<i>") and segment.endswith("</i>"):
            _flush(segment[3:-4], italic=True)
        elif segment.startswith("<sup>") and segment.endswith("</sup>"):
            run = para.add_run(segment[5:-6])
            run.font.superscript = True
        elif segment.startswith("<math>") and segment.endswith("</math>"):
            _flush(segment[6:-7])   # testo normale, tag rimossi
        else:
            _flush(segment)


# ---------------------------------------------------------------------------
# Funzioni per l'export strutturato Surya 0.20
# ---------------------------------------------------------------------------

def _html_to_markup(html: str) -> str:
    """Converte l'HTML di un blocco Surya nel markup semplificato (<i>, <sup>)."""
    html = re.sub(r'<em[^>]*>', '<i>', html, flags=re.IGNORECASE)
    html = re.sub(r'</em>', '</i>', html, flags=re.IGNORECASE)
    # Grassetto: rimuovi i tag wrappanti (nessun grassetto nel markup)
    html = re.sub(r'<b[^>]*>(.*?)</b>', r'\1', html, flags=re.IGNORECASE | re.DOTALL)
    html = re.sub(r'<strong[^>]*>(.*?)</strong>', r'\1', html, flags=re.IGNORECASE | re.DOTALL)
    # Interruzioni di riga e separatori paragrafo
    html = re.sub(r'<br\s*/?>', '\n', html, flags=re.IGNORECASE)
    html = re.sub(r'</p>\s*<p[^>]*>', '\n', html, flags=re.IGNORECASE)
    html = re.sub(r'<p[^>]*>|</p>', '', html, flags=re.IGNORECASE)
    # Liste: converti <li> in righe
    html = re.sub(r'<li[^>]*>', '• ', html, flags=re.IGNORECASE)
    html = re.sub(r'</li>', '\n', html, flags=re.IGNORECASE)
    html = re.sub(r'<[ou]l[^>]*>|</[ou]l>', '', html, flags=re.IGNORECASE)
    # Math: mantieni solo il contenuto
    html = re.sub(r'<math[^>]*>(.*?)</math>', r'\1', html, flags=re.IGNORECASE | re.DOTALL)
    # Preserva <i>, </i>, <sup>N</sup> tramite placeholder prima di strippare
    html = html.replace('<i>', '\x00ITAG\x00')
    html = html.replace('</i>', '\x00EITAG\x00')
    html = re.sub(r'<sup>(\d+)</sup>', '\x00SUP\\1\x00', html)
    html = re.sub(r'<[^>]+>', '', html)
    html = html.replace('\x00ITAG\x00', '<i>')
    html = html.replace('\x00EITAG\x00', '</i>')
    html = re.sub(r'\x00SUP(\d+)\x00', r'<sup>\1</sup>', html)
    return html.strip()


class _TableParser(HTMLParser):
    """Parser minimale per estrarre righe e celle da un frammento HTML di tabella."""

    def __init__(self):
        super().__init__()
        self.rows = []
        self._in_row = False
        self._in_cell = False
        self._cell_buf = []

    def handle_starttag(self, tag, attrs):
        if tag == "tr":
            self._in_row = True
            self.rows.append([])
        elif tag in ("td", "th") and self._in_row:
            self._in_cell = True
            self._cell_buf = []

    def handle_endtag(self, tag):
        if tag in ("td", "th") and self._in_cell:
            self.rows[-1].append("".join(self._cell_buf).strip())
            self._in_cell = False
            self._cell_buf = []
        elif tag == "tr":
            self._in_row = False

    def handle_data(self, data):
        if self._in_cell:
            self._cell_buf.append(data)


def _add_html_table(doc, html: str) -> None:
    """Aggiunge una tabella Word da un frammento HTML <table>."""
    parser = _TableParser()
    parser.feed(html)
    rows = [r for r in parser.rows if r]
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    if ncols == 0:
        return
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            if c_idx < ncols:
                table.rows[r_idx].cells[c_idx].text = cell_text


def _add_block_to_docx(doc, block: dict) -> None:
    """Aggiunge un blocco Surya 0.20 al documento Word."""
    label = block.get("label", "Text")
    html = block.get("html", "")

    if label in _SKIP_DOCX_LABELS:
        return

    if label in _IMAGE_PLACEHOLDER:
        # Blocco immagine: nessun testo, ma resta in reading order come
        # marcatore così l'utente si orienta nel layout originale.
        doc.add_paragraph(style="Caption").add_run(
            _(_IMAGE_PLACEHOLDER[label])).italic = True
        return

    if not html:
        return

    if label == "Table":
        _add_html_table(doc, html)
        return

    markup = _html_to_markup(html)
    if not markup:
        return

    if label == "ListGroup":
        # Un blocco lista contiene più voci separate da newline (dai <li>).
        for line in markup.split("\n"):
            item = line.lstrip("• ").strip()
            if item:
                _add_markup_runs(doc.add_paragraph(style="List Bullet"), item)
        return

    if label == "Footnote":
        para = doc.add_paragraph(style="Normal")
        para.add_run("— ").italic = True
    else:
        style = _BLOCK_STYLE.get(label, "Normal")
        para = doc.add_paragraph(style=style)

    _add_markup_runs(para, markup)


# ---------------------------------------------------------------------------
# Funzioni di scrittura file
# ---------------------------------------------------------------------------

def write_txt(text: str, path: str) -> None:
    """Salva il testo come file UTF-8, convertendo i tag in Unicode."""
    with open(path, "w", encoding="utf-8") as f:
        f.write(_convert_markup(text))


def write_docx(text: str, path: str, blocks=None) -> None:
    """Salva il testo come .docx.

    Se blocks è fornito (list[list[dict]] da Surya 0.20), usa il rendering
    strutturato con heading, tabelle e stili semantici. Altrimenti usa il
    rendering flat con corsivo, apici e interruzioni di pagina Word nativi.
    """
    doc = docx.Document()
    if blocks:
        for page_idx, page_blocks in enumerate(blocks):
            if page_idx > 0:
                doc.add_page_break()
            for block in page_blocks:
                _add_block_to_docx(doc, block)
    else:
        pages = text.split("\f")
        for page_idx, page_text in enumerate(pages):
            if page_idx > 0:
                doc.add_page_break()
            for para_text in page_text.split("\n\n"):
                para_text = para_text.strip()
                if not para_text:
                    continue
                para = doc.add_paragraph()
                _add_markup_runs(para, para_text)
    doc.save(path)


def write_html(html: str, path: str) -> None:
    """Salva l'HTML strutturato Surya 0.20 come file .html con wrapper CSS."""
    content = (
        '<!DOCTYPE html>\n<html lang="it">\n<head>\n'
        '<meta charset="utf-8">\n'
        '<title>OCR Lab — risultato</title>\n'
        '<style>\n'
        '  body { font-family: Georgia, serif; max-width: 900px; margin: 2em auto;'
        ' padding: 0 1em; line-height: 1.5; }\n'
        '  .title { font-size: 1.6em; font-weight: bold; margin: 0.8em 0 0.4em; }\n'
        '  .section-header { font-size: 1.25em; font-weight: bold; margin: 1em 0 0.4em; }\n'
        '  .caption { font-size: 0.9em; font-style: italic; color: #555; }\n'
        '  .footnote { font-size: 0.85em; border-top: 1px solid #ccc;'
        ' margin-top: 1.5em; padding-top: 0.5em; color: #444; }\n'
        '  .page-header, .page-footer { font-size: 0.8em; color: #888; }\n'
        '  .equation { font-style: italic; }\n'
        '  table { border-collapse: collapse; width: 100%; margin: 1em 0; }\n'
        '  td, th { border: 1px solid #bbb; padding: 4px 8px; }\n'
        '  th { background: #f0f0f0; font-weight: bold; }\n'
        '  hr.page-break { border: none; border-top: 2px dashed #aaa; margin: 2em 0; }\n'
        '</style>\n</head>\n<body>\n'
        + html
        + '\n</body>\n</html>'
    )
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)


# Font per il testo OCR invisibile. Serve copertura Unicode piena: i base-14 di
# PyMuPDF ("helv"…) usano codifica Latin-1 e trasformano €, trattini tipografici,
# virgolette curve e ellissi in '?'. Un TrueType di sistema li mappa (con ToUnicode)
# → il testo estratto/letto dagli screen reader resta fedele.
#
# Il testo OCR viene disegnato con TextWriter (vedi _write_text_layer), NON con
# insert_textbox: quest'ultimo passa per lo shaping OpenType, che altera il testo
# *estratto* (parentesi → ﴾ ﴿ con segoeui, glifi mancanti → null, ecc.). TextWriter
# invece costruisce il ToUnicode dalla stringa di input, quindi il testo estratto
# resta fedele a prescindere dai glifi (che qui sono invisibili). Verificato con
# PyMuPDF 1.27 estraendo dal PDF prodotto: greco politonico compreso.
#
# Il font conta ancora per un dettaglio: lo spazio. Con arial/tahoma/calibri/times/
# verdana/cambria/palatino lo spazio U+0020 viene estratto come nbsp U+00A0
# (rompe la ricerca Ctrl+F di frasi); seguisym.ttf (Segoe UI Symbol) lo preserva
# e ha copertura larghissima → scelta migliore su Windows.
_OCR_FONT_CANDIDATES = {
    "win32": [
        r"C:\Windows\Fonts\seguisym.ttf",  # spazi + parentesi puliti, copertura larga
        r"C:\Windows\Fonts\arial.ttf",     # fallback: nbsp al posto dello spazio
        r"C:\Windows\Fonts\tahoma.ttf",
    ],
    "darwin": [
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial.ttf",
        "/System/Library/Fonts/Helvetica.ttc",
    ],
}

_ocr_font_cache = None


def _resolve_ocr_font() -> tuple[str, str | None]:
    """Ritorna (fontname, fontfile) per il testo OCR invisibile.

    Preferisce un TrueType di sistema a copertura Unicode piena; fallback su
    'helv' (base-14, solo Latin-1) se nessun candidato esiste sul sistema.
    Il risultato è memorizzato in cache: la risoluzione avviene una sola volta.
    """
    global _ocr_font_cache
    if _ocr_font_cache is not None:
        return _ocr_font_cache
    for path in _OCR_FONT_CANDIDATES.get(sys.platform, []):
        if os.path.exists(path):
            _ocr_font_cache = ("ocrfont", path)
            return _ocr_font_cache
    _ocr_font_cache = ("helv", None)
    return _ocr_font_cache


def _write_text_layer(fitz, doc, page, rect, text: str, font) -> int:
    """Disegna il testo OCR invisibile (render_mode=3) con TextWriter.

    TextWriter (invece di insert_textbox) preserva il testo estratto per qualsiasi
    carattere, ma non fa né word-wrap né auto-shrink: li implementiamo qui.
    Va a capo sulle parole per stare nella larghezza e sceglie la fontsize più
    grande (≤6pt) con cui tutte le righe entrano in altezza, così su pagine dense
    non si perde testo (parità con il vecchio _fit_fontsize).

    Percorso di ripiego (motori senza bbox): tutto il testo di pagina finisce in
    un unico content stream marcato `/P <</MCID 0>>`. Ritorna il numero di MCID
    scritti (0 o 1), coerente con _write_positioned_text_layer.
    """
    margin  = 10
    max_w   = rect.width  - 2 * margin
    avail_h = rect.height - 2 * margin
    paragraphs = text.split("\n")

    def wrap(fs):
        lines = []
        for para in paragraphs:
            para = para.rstrip()
            if not para:
                lines.append("")          # preserva lo stacco fra paragrafi
                continue
            cur = ""
            for word in para.split(" "):
                trial = word if not cur else cur + " " + word
                if not cur or font.text_length(trial, fontsize=fs) <= max_w:
                    cur = trial
                else:
                    lines.append(cur)
                    cur = word
            lines.append(cur)
        return lines

    fs, lines = 2, wrap(2)
    for cand in (6, 5, 4, 3, 2):
        wl = wrap(cand)
        if len(wl) * (cand * 1.2) <= avail_h:
            fs, lines = cand, wl
            break

    if not any(ln.strip() for ln in lines):
        return 0

    line_h = fs * 1.2
    tw = fitz.TextWriter(rect)
    x = rect.x0 + margin
    y = rect.y0 + margin + fs
    wrote = False
    for ln in lines:
        if ln:
            tw.append((x, y), ln, font=font, fontsize=fs)
            wrote = True
        y += line_h
    if not wrote:
        return 0
    tw.write_text(page, render_mode=3)
    xref = page.get_contents()[-1]
    _wrap_marked_content(doc=doc, xref=xref, prefix=b"/P <</MCID 0>> BDC")
    return 1


def _block_plain_text(html: str) -> str:
    """Testo piano di un blocco Surya: rimuove i tag, decodifica le entità e
    normalizza gli spazi (per il layer OCR invisibile posizionato)."""
    text = re.sub(r"<br\s*/?>", " ", html, flags=re.IGNORECASE)
    text = re.sub(r"<[^>]+>", " ", text)
    text = unescape(text)
    return re.sub(r"\s+", " ", text).strip()


def _wrap_words(font, words: list, fs: float, max_w: float) -> list:
    """Word-wrap di una lista di parole alla larghezza max_w per fontsize fs."""
    lines, cur = [], ""
    for word in words:
        trial = word if not cur else cur + " " + word
        if not cur or font.text_length(trial, fontsize=fs) <= max_w:
            cur = trial
        else:
            lines.append(cur)
            cur = word
    if cur:
        lines.append(cur)
    return lines


def _write_positioned_text_layer(fitz, doc, page, rect, page_blocks, font) -> int:
    """Disegna il testo OCR invisibile posizionato sulle bbox dei blocchi.

    Ogni blocco con bbox (normalizzata in [0,1], vedi surya20_worker per Surya o
    _data_to_text_and_blocks per Tesseract/Windows OCR) viene riportato nel
    rettangolo corrispondente della pagina, scegliendo la fontsize più grande con
    cui il testo (word-wrappato alla larghezza del blocco) sta nell'altezza del
    blocco. Così il layer invisibile combacia con la posizione del testo
    nell'immagine: selezione ed evidenziazione cadono al posto giusto.

    **Un content stream + un MCID per blocco** (non più uno unico per pagina):
    ogni blocco diventa un `/P <</MCID n>>` a sé, così lo structure tree espone i
    blocchi come elementi distinti (paragrafo per Surya, riga per Tesseract/WinOCR)
    e la lettura riga-per-riga dello screen reader segue la topologia della pagina,
    come nei PDF-scansione taggati di riferimento. Ritorna il numero di MCID
    scritti (= blocchi con testo), che _add_tag_structure userà per creare uno
    StructElem per blocco.
    """
    W, H = rect.width, rect.height
    mcid = 0
    for blk in page_blocks:
        bbox = blk.get("bbox")
        if not bbox:
            continue
        text = _block_plain_text(blk.get("html", ""))
        if not text:
            continue
        x1, y1, x2, y2 = bbox
        box_x0 = rect.x0 + x1 * W
        box_y0 = rect.y0 + y1 * H
        box_w = max(1.0, (x2 - x1) * W)
        box_h = max(1.0, (y2 - y1) * H)

        words = text.split(" ")
        size = min(box_h, 60.0)
        lines = _wrap_words(font, words, size, box_w)
        while size > 1.0 and len(lines) * size * 1.2 > box_h:
            size -= 0.5
            lines = _wrap_words(font, words, size, box_w)

        line_h = size * 1.2
        y = box_y0 + size
        tw = fitz.TextWriter(rect)
        wrote = False
        for ln in lines:
            if ln:
                # Spazio finale reale: i lettori che leggono il testo taggato
                # (NVDA, "copia" da PDF taggato) concatenano i glifi nell'ordine
                # del content stream, ignorando la geometria. Senza separatore i
                # run adiacenti si fondono ("RECENSIONICRITICHE", "delladesuetudine");
                # lo spazio a fine riga/blocco mantiene il testo leggibile.
                tw.append((box_x0, y), ln + " ", font=font, fontsize=size)
                wrote = True
            y += line_h
        if wrote:
            tw.write_text(page, render_mode=3)
            xref = page.get_contents()[-1]
            _wrap_marked_content(doc=doc, xref=xref,
                                 prefix=f"/P <</MCID {mcid}>> BDC".encode())
            mcid += 1
    return mcid


def _wrap_marked_content(doc, xref: int, prefix: bytes) -> None:
    """Avvolge il content stream `xref` in una sequenza di marked content.

    Antepone `prefix` (es. b"/Artifact BMC" o b"/P <</MCID 0>> BDC") e accoda
    `EMC`, così il disegno esistente (immagine o testo) diventa un blocco
    marcato riconoscibile dallo structure tree / dalle tecnologie assistive.
    """
    body = doc.xref_stream(xref)
    doc.update_stream(xref, prefix + b"\n" + body + b"\nEMC\n")


def _add_tag_structure(doc, page_tags: list) -> None:
    """Costruisce lo structure tree (tagged PDF) del documento.

    page_tags: lista di (page_xref, n_mcids) — un elemento per pagina, in ordine.
    Ogni pagina con testo ha nel suo content stream `n_mcids` blocchi marcati
    `/P <</MCID k>> BDC … EMC` (k = 0..n_mcids-1; l'immagine è marcata `/Artifact`,
    quindi fuori dall'albero). Per ogni MCID creiamo un StructElem /P distinto, così
    lo screen reader espone i blocchi come elementi separati (paragrafo/riga) e la
    lettura riga-per-riga segue la topologia della pagina.

    Replica lo schema dei PDF taggati da Word: StructTreeRoot → /Document →
    [ /P, /P, … ], con ParentTree e /StructParents per pagina. La voce del
    ParentTree di ogni pagina è un array indicizzato per MCID → StructElem.
    """
    with_text = [(px, n) for (px, n) in page_tags if n > 0]
    if not with_text:
        return

    catalog_xref = doc.pdf_catalog()
    root_xref    = doc.get_new_xref()   # StructTreeRoot
    doc_xref     = doc.get_new_xref()   # elemento /Document

    all_p_xrefs = []   # tutti gli StructElem /P, in ordine di lettura (per /Document)
    num_parts   = []   # voci del ParentTree: structParentIndex → [ /P per MCID ]
    for struct_parent, (page_xref, n_mcids) in enumerate(with_text):
        page_p_xrefs = []
        for mcid in range(n_mcids):
            p_xref = doc.get_new_xref()
            doc.update_object(p_xref, (
                f"<< /Type /StructElem /S /P "
                f"/P {doc_xref} 0 R /Pg {page_xref} 0 R /K [{mcid}] >>"
            ))
            page_p_xrefs.append(p_xref)
        all_p_xrefs.extend(page_p_xrefs)

        # Array del ParentTree per la pagina: indice = MCID, valore = StructElem.
        arr_xref = doc.get_new_xref()
        doc.update_object(arr_xref,
                          "[ " + " ".join(f"{px} 0 R" for px in page_p_xrefs) + " ]")
        num_parts.append(f"{struct_parent} {arr_xref} 0 R")

        doc.xref_set_key(page_xref, "StructParents", str(struct_parent))
        doc.xref_set_key(page_xref, "Tabs", "/S")

    kids = " ".join(f"{px} 0 R" for px in all_p_xrefs)
    doc.update_object(doc_xref, (
        f"<< /Type /StructElem /S /Document "
        f"/P {root_xref} 0 R /K [ {kids} ] >>"
    ))

    parent_tree_xref = doc.get_new_xref()
    doc.update_object(parent_tree_xref,
                      f"<< /Nums [ {' '.join(num_parts)} ] >>")

    doc.update_object(root_xref, (
        f"<< /Type /StructTreeRoot /K [ {doc_xref} 0 R ] "
        f"/ParentTree {parent_tree_xref} 0 R /ParentTreeNextKey {len(with_text)} >>"
    ))

    doc.xref_set_key(catalog_xref, "StructTreeRoot", f"{root_xref} 0 R")
    doc.xref_set_key(catalog_xref, "MarkInfo", "<< /Marked true >>")


def _detect_rotation(pix) -> int:
    """Gradi (0/90/180/270) di cui ruotare l'immagine in senso ORARIO per
    raddrizzarla, via Tesseract OSD. Ritorna 0 se non rilevabile o poco
    affidabile (Tesseract assente/non configurato, pagina senza testo, bassa
    confidenza), lasciando la pagina invariata.

    Self-contained: opera sul pixmap già rasterizzato, quindi vale per qualsiasi
    motore (Surya, Tesseract, Windows OCR, …) senza toccare il codice OCR. La
    dipendenza da Tesseract è soft (guardata): se manca, niente rotazione.
    """
    try:
        import io as _io
        import pytesseract
        from PIL import Image
        from app.engine.ocr_engine import _configure_tesseract
        _configure_tesseract()
        img = Image.open(_io.BytesIO(pix.tobytes("png")))
        osd = pytesseract.image_to_osd(img, output_type=pytesseract.Output.DICT)
        rot = int(osd.get("rotate", 0)) % 360
        conf = float(osd.get("orientation_conf", 0) or 0)
        # Agiamo solo su rotazioni cardinali con un minimo di confidenza: una
        # rotazione errata su pagina già dritta sarebbe peggio dello stato attuale.
        if rot in (90, 180, 270) and conf >= 1.0:
            return rot
    except Exception:
        pass
    return 0


def _rotate_pixmap(fitz, pix, rot: int):
    """Ruota il pixmap di `rot` gradi in senso orario, ritornando un nuovo pixmap
    raddrizzato (via PIL). `rot` deve essere 90/180/270."""
    import io as _io
    from PIL import Image
    img = Image.open(_io.BytesIO(pix.tobytes("png")))
    # PIL.rotate ruota in senso antiorario: -rot per ottenere l'orario.
    img = img.rotate(-rot, expand=True)
    buf = _io.BytesIO()
    img.save(buf, format="png")
    return fitz.Pixmap(buf.getvalue())


def _rotate_bbox(bbox, rot: int):
    """Trasforma una bbox normalizzata [x1,y1,x2,y2] (origine in alto a sinistra)
    quando l'immagine viene ruotata di `rot` gradi in senso orario. Gli assi si
    scambiano per 90/270; i valori restano in [0,1] rispetto alle NUOVE dimensioni.
    """
    x1, y1, x2, y2 = bbox
    if rot == 90:
        pts = [(1 - y1, x1), (1 - y2, x2)]
    elif rot == 180:
        pts = [(1 - x1, 1 - y1), (1 - x2, 1 - y2)]
    elif rot == 270:
        pts = [(y1, 1 - x1), (y2, 1 - x2)]
    else:
        return [x1, y1, x2, y2]
    xs = [p[0] for p in pts]
    ys = [p[1] for p in pts]
    return [min(xs), min(ys), max(xs), max(ys)]


def write_searchable_pdf(source_path: str, ocr_text: str, out_path: str,
                         blocks=None, line_blocks=None, angles=None) -> None:
    """PDF ricercabile e **taggato**: immagine di sfondo + testo OCR leggibile.

    Struttura di ogni pagina:
    1. Immagine rasterizzata sorgente, marcata `/Artifact`: è puro sfondo
       visivo e le tecnologie assistive la ignorano (niente "grafico" letto).
       Se la scansione è ruotata (di lato), viene raddrizzata via OSD prima di
       essere inserita (vedi _detect_rotation/_rotate_pixmap).
    2. Testo OCR invisibile (render_mode=3), un blocco marcato `/P <</MCID k>>
       BDC … EMC` **per ogni blocco/riga** (non più uno unico per pagina), con un
       font a copertura Unicode piena (vedi _resolve_ocr_font).

    Sopra le pagine viene costruito uno structure tree completo (vedi
    _add_tag_structure): StructTreeRoot + MarkInfo/Marked. È lo stesso schema
    dei PDF taggati da Word — l'unico che NVDA/Acrobat leggono in modo pulito
    (testo una sola volta, nell'ordine giusto, senza doppioni).

    Il tentativo precedente di taggare falliva perché lo StructElem puntava a
    un MCID che non esisteva nel content stream (insert_textbox non emette
    marked content): qui il BDC/EMC viene scritto esplicitamente, così il tag
    e il glyph combaciano.

    Se `blocks` (list[list[dict]] semantici da Surya 0.20, con bbox normalizzate)
    o `line_blocks` (blocchi posizionali riga-per-riga da Tesseract/Windows OCR,
    stesso schema {label, html, bbox}) sono forniti, il testo di ogni pagina viene
    posizionato sopra la regione corrispondente dell'immagine (vedi
    _write_positioned_text_layer): selezione/evidenziazione e lettura spaziale
    combaciano col layout, riga dopo riga, come nei PDF-scansione taggati (es.
    ABBYY). Ogni blocco/riga è un `/P` a sé nello structure tree, quindi la
    lettura riga-per-riga dello screen reader segue la topologia della pagina.
    Senza bbox si ripiega sul layer a flusso (_write_text_layer), con un unico
    `/P` per pagina.

    `blocks` e `line_blocks` si escludono a vicenda per una data pagina: i primi
    portano anche la semantica (usata dal DOCX), i secondi solo la geometria. Qui
    conta solo la geometria, quindi usiamo `blocks or line_blocks`.

    Args:
        source_path: percorso del file sorgente (PDF o immagine).
        ocr_text:    testo OCR con pagine separate da \\f.
        out_path:    percorso del file PDF di output.
        blocks:      opzionale, list[list[dict]] semantici per pagina (Surya).
        line_blocks: opzionale, list[list[dict]] posizionali riga-per-riga
                     (Tesseract/Windows OCR). Ignorato se `blocks` è presente.
    """
    try:
        import pymupdf as fitz
    except ImportError:
        import fitz

    blocks = blocks or line_blocks

    pages_text = [strip_markup(t).strip() for t in ocr_text.split("\f")]

    ext = os.path.splitext(source_path)[1].lower()
    if ext == ".pdf":
        src = fitz.open(source_path)
    else:
        img_doc = fitz.open(source_path)
        pdf_bytes = img_doc.convert_to_pdf()
        img_doc.close()
        src = fitz.open("pdf", pdf_bytes)

    out_doc = fitz.open()
    fontname, fontfile = _resolve_ocr_font()
    ocr_font = fitz.Font(fontfile=fontfile) if fontfile else fitz.Font(fontname=fontname)
    page_tags = []   # (page_xref, n_mcids) per _add_tag_structure

    try:
        for i in range(len(src)):
            src_page = src[i]
            src_rect = src_page.rect   # dimensioni pagina in punti

            # Rasterizza la pagina sorgente e, se ruotata (scansione di lato),
            # raddrizzala: l'immagine di sfondo e le bbox del testo vengono ruotate
            # insieme, così l'output è dritto come l'originale e la lettura
            # riga-per-riga segue la topologia (parità coi PDF-scansione taggati).
            pix = src_page.get_pixmap(dpi=200)
            # Se l'OCR ha già raddrizzato la pagina (angolo fornito dall'engine, es.
            # Surya 0.20 sugli spread ruotati), le sue bbox sono GIÀ dritte: usiamo
            # quell'angolo per l'immagine e NON ri-ruotiamo le bbox (evita la doppia
            # rotazione). Altrimenti autorileviamo via OSD e ruotiamo anche le bbox.
            ocr_rot = angles[i] if angles and i < len(angles) else None
            if ocr_rot:
                rot = ocr_rot
                rotate_bboxes = False
            else:
                rot = _detect_rotation(pix)
                rotate_bboxes = True
            if rot:
                pix = _rotate_pixmap(fitz, pix, rot)

            # La pagina resta dimensionata in PUNTI; su 90/270 gli assi si scambiano.
            if rot in (90, 270):
                page_w, page_h = src_rect.height, src_rect.width
            else:
                page_w, page_h = src_rect.width, src_rect.height
            rect = fitz.Rect(0, 0, page_w, page_h)

            out_page = out_doc.new_page(width=page_w, height=page_h)

            # 1. Immagine sorgente come sfondo visivo, marcata /Artifact.
            out_page.insert_image(rect, pixmap=pix)
            img_xref = out_page.get_contents()[-1]
            _wrap_marked_content(doc=out_doc, xref=img_xref, prefix=b"/Artifact BMC")

            # 2. Testo OCR invisibile (render_mode=3): un /P con MCID per blocco.
            #    Se la pagina ha blocchi con bbox li posizioniamo sull'immagine
            #    (ruotando le bbox come l'immagine), altrimenti layer a flusso.
            text = pages_text[i] if i < len(pages_text) else ""
            page_blocks = blocks[i] if blocks and i < len(blocks) else None
            has_boxes = bool(page_blocks) and any(b.get("bbox") for b in page_blocks)
            n_mcids = 0
            if has_boxes:
                if rot and rotate_bboxes:
                    page_blocks = [
                        {**b, "bbox": _rotate_bbox(b["bbox"], rot)} if b.get("bbox")
                        else b
                        for b in page_blocks
                    ]
                n_mcids = _write_positioned_text_layer(
                    fitz, out_doc, out_page, rect, page_blocks, ocr_font)
                if n_mcids == 0 and text:
                    n_mcids = _write_text_layer(
                        fitz, out_doc, out_page, rect, text, ocr_font)
            elif text:
                n_mcids = _write_text_layer(
                    fitz, out_doc, out_page, rect, text, ocr_font)

            page_tags.append((out_page.xref, n_mcids))

        _add_tag_structure(out_doc, page_tags)
        out_doc.save(out_path, garbage=4, deflate=True)
    finally:
        src.close()
        out_doc.close()


def write_file(text: str, path: str, source_path: str = "",
               blocks=None, html: str = "", line_blocks=None,
               angles=None) -> None:
    """Salva il testo in base all'estensione del file.

    blocks:      list[list[dict]] semantici da Surya 0.20 (layout strutturato per
                 il DOCX e posizionamento del testo sul PDF ricercabile).
    html:        HTML grezzo da Surya 0.20 (necessario per l'esportazione .html).
    line_blocks: list[list[dict]] posizionali riga-per-riga da Tesseract/Windows
                 OCR. Servono solo al PDF ricercabile (topologia testo/immagine);
                 il DOCX li ignora e resta sul rendering a flusso, migliore per
                 blocchi privi di semantica.
    angles:      list[int|None] per pagina: angolo ORARIO con cui l'OCR ha già
                 raddrizzato la pagina (bbox dei blocchi già dritte). Il PDF
                 ricercabile raddrizza l'immagine con questo angolo SENZA
                 ri-ruotare le bbox. None su una pagina → il writer si autorileva.
    """
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        write_docx(text, path, blocks=blocks)
    elif ext == ".txt":
        write_txt(text, path)
    elif ext == ".pdf":
        if not source_path:
            raise ValueError(
                "Per il PDF ricercabile è necessario il percorso del file sorgente."
            )
        write_searchable_pdf(source_path, text, path, blocks=blocks,
                             line_blocks=line_blocks, angles=angles)
    elif ext == ".html":
        if not html:
            raise ValueError(
                "L'esportazione HTML è disponibile solo con il motore Surya 0.2."
            )
        write_html(html, path)
    else:
        raise ValueError(f"Formato non supportato: {ext}")
